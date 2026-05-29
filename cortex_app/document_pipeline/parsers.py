import base64
import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path

import litellm

from config.settings import get_settings

settings = get_settings()


@dataclass
class ParsedChunkRaw:
    text: str
    page_start: int
    page_end: int
    section: str | None
    chunk_type: str  # "text" | "table"


def _image_exts() -> set[str]:
    return {e for e in settings.KB_SUPPORTED_EXTENSIONS if e in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}}


def parse_document(file_bytes: bytes, filename: str, openai_api_key: str) -> list[ParsedChunkRaw]:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    if ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    if ext in (".csv", ".xlsx", ".xls"):
        return _parse_tabular(file_bytes, ext)
    if ext in (".txt", ".md"):
        return _parse_text(file_bytes)
    if ext in _image_exts():
        import asyncio
        return asyncio.run(_parse_image(file_bytes, filename, openai_api_key))
    raise ValueError(f"Unsupported extension: {ext}")


def compute_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(content: bytes) -> list[ParsedChunkRaw]:
    import pdfplumber

    chunks: list[ParsedChunkRaw] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()
            table_bboxes = [t.bbox for t in page.find_tables()] if tables else []

            for table_data in tables:
                if not table_data:
                    continue
                markdown = _table_to_markdown(table_data)
                if markdown.strip():
                    chunks.append(ParsedChunkRaw(
                        text=markdown,
                        page_start=page_num,
                        page_end=page_num,
                        section=None,
                        chunk_type="table",
                    ))

            text = page.filter(
                lambda obj: not _is_in_bbox(obj, table_bboxes)
            ).extract_text() or ""
            text = _clean_text(text)
            if text.strip():
                chunks.append(ParsedChunkRaw(
                    text=text,
                    page_start=page_num,
                    page_end=page_num,
                    section=_extract_section_heading(text),
                    chunk_type="text",
                ))
    return chunks


def _table_to_markdown(table_data: list[list]) -> str:
    if not table_data:
        return ""
    rows = []
    for i, row in enumerate(table_data):
        cells = [str(c or "").strip() for c in row]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


def _is_in_bbox(obj, bboxes: list) -> bool:
    if not hasattr(obj, "bbox"):
        return False
    ox0, oy0, ox1, oy1 = obj["bbox"] if isinstance(obj, dict) else obj.bbox
    for bx0, by0, bx1, by1 in bboxes:
        if ox0 >= bx0 and oy0 >= by0 and ox1 <= bx1 and oy1 <= by1:
            return True
    return False


def _clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()


def _extract_section_heading(text: str) -> str | None:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) >= 2 and len(lines[0]) < 100 and len(lines[1]) > 50:
        return lines[0]
    return None


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(content: bytes) -> list[ParsedChunkRaw]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    chunks: list[ParsedChunkRaw] = []
    current_section: str | None = None
    current_text_parts: list[str] = []
    paragraph_index = 0

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]

        if tag == "p":
            para = doc.paragraphs[paragraph_index] if paragraph_index < len(doc.paragraphs) else None
            paragraph_index += 1
            if not para:
                continue

            style = para.style.name if para.style else ""
            text = para.text.strip()

            if style.startswith("Heading"):
                if current_text_parts:
                    chunks.append(ParsedChunkRaw(
                        text="\n".join(current_text_parts),
                        page_start=1, page_end=1,
                        section=current_section,
                        chunk_type="text",
                    ))
                    current_text_parts = []
                current_section = text
            elif text:
                current_text_parts.append(text)

        elif tag == "tbl":
            if current_text_parts:
                chunks.append(ParsedChunkRaw(
                    text="\n".join(current_text_parts),
                    page_start=1, page_end=1,
                    section=current_section,
                    chunk_type="text",
                ))
                current_text_parts = []

            markdown = _docx_table_to_markdown(block, doc)
            if markdown.strip():
                chunks.append(ParsedChunkRaw(
                    text=markdown,
                    page_start=1, page_end=1,
                    section=current_section,
                    chunk_type="table",
                ))

    if current_text_parts:
        chunks.append(ParsedChunkRaw(
            text="\n".join(current_text_parts),
            page_start=1, page_end=1,
            section=current_section,
            chunk_type="text",
        ))

    return chunks


def _docx_table_to_markdown(tbl_element, doc) -> str:
    from docx.oxml.ns import qn
    rows = []
    tr_elements = tbl_element.findall(f".//{qn('w:tr')}")
    for i, tr in enumerate(tr_elements):
        cells = [tc.text_content() if hasattr(tc, "text_content") else
                 "".join(r.text or "" for r in tr.iter())
                 for tc in tr.findall(f".//{qn('w:tc')}")]
        cells = [c.strip() for c in cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


# ---------------------------------------------------------------------------
# CSV / XLSX
# ---------------------------------------------------------------------------

def _parse_tabular(content: bytes, ext: str) -> list[ParsedChunkRaw]:
    import pandas as pd

    if ext == ".csv":
        df = pd.read_csv(io.BytesIO(content))
    else:
        df = pd.read_excel(io.BytesIO(content))

    total_rows, total_cols = df.shape
    chunks: list[ParsedChunkRaw] = []

    meta_cols = ", ".join(str(c) for c in df.columns[:30])
    if total_cols > 30:
        meta_cols += f" (+{total_cols - 30} more columns)"
    meta = f"File: {total_rows} rows × {total_cols} columns\nColumns: {meta_cols}"
    chunks.append(ParsedChunkRaw(
        text=meta, page_start=1, page_end=1,
        section="Metadata", chunk_type="table",
    ))

    rows_per_chunk = settings.KB_CSV_ROWS_PER_CHUNK
    for i in range(0, len(df), rows_per_chunk):
        group = df.iloc[i: i + rows_per_chunk]
        try:
            markdown = group.to_markdown(index=False)
        except Exception:
            markdown = group.to_string(index=False)
        chunks.append(ParsedChunkRaw(
            text=markdown,
            page_start=1, page_end=1,
            section=f"Rows {i + 1}–{i + len(group)}",
            chunk_type="table",
        ))

    return chunks


# ---------------------------------------------------------------------------
# TXT / MD
# ---------------------------------------------------------------------------

def _parse_text(content: bytes) -> list[ParsedChunkRaw]:
    text = content.decode("utf-8", errors="replace")
    return [ParsedChunkRaw(
        text=_clean_text(text),
        page_start=1, page_end=1,
        section=None,
        chunk_type="text",
    )]


# ---------------------------------------------------------------------------
# Image (vision)
# ---------------------------------------------------------------------------

async def _parse_image(content: bytes, filename: str, openai_api_key: str) -> list[ParsedChunkRaw]:
    ext = Path(filename).suffix.lower().lstrip(".")
    mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "webp": "webp", "gif": "gif", "bmp": "bmp"}
    mime = f"image/{mime_map.get(ext, ext)}"
    b64 = base64.b64encode(content).decode()

    response = await litellm.acompletion(
        model=settings.VISION_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "Describe the content of this image in detail, including any text, charts, diagrams, or visual information present."},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ],
        }],
        api_key=openai_api_key,
    )
    description = response.choices[0].message.content or ""
    return [ParsedChunkRaw(
        text=description,
        page_start=1, page_end=1,
        section=None,
        chunk_type="text",
    )]
